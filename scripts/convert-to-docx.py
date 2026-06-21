#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将深夜小茶馆风格的 .md 故事文件转换为 .docx Word 文档。

用法：
    python convert-to-docx.py <输入.md> [输出.docx]

如果不指定输出路径，默认与输入同目录、同文件名、.docx 后缀。
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要 python-docx 库。请运行: pip install python-docx")
    sys.exit(1)


def set_run_font(run, font_name="微软雅黑", size=None, bold=False, color=None):
    """设置 run 的字体属性（中文字体 + 西文字体）。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_title(doc, text):
    """添加故事标题——居中、大号、加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=22, bold=True)


def add_section_header(doc, text):
    """添加 #01 #02 等章节标题——居中、加粗、中型字。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)


def add_paragraph(doc, text):
    """添加普通叙述段落。支持行内 **bold** 标记。"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 两字符缩进
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5

    # 解析行内 **bold**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=12, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run, size=12)


def add_blockquote(doc, text):
    """添加引用块（如结尾的 > 奇闻说今古……）。"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    # 去掉行首的 > 和空格
    cleaned = re.sub(r'^>\s?', '', text)
    # 解析 bold
    parts = re.split(r'(\*\*[^*]+\*\*)', cleaned)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, size=11, bold=True, color=(128, 128, 128))
        else:
            run = p.add_run(part)
            set_run_font(run, size=11, color=(128, 128, 128))


def add_end_marker(doc):
    """添加 END 标记——居中。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("END")
    set_run_font(run, size=14, bold=True, color=(150, 150, 150))


def convert_md_to_docx(md_path: Path, docx_path: Path):
    """将 markdown 故事文件转换为 Word 文档。"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # 设置默认字体 + 默认段间距（去掉段与段之间的空行）
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(12)
    # 段落之间用 6pt 间距代替空行
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    i = 0

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # 跳过空行——用段间距代替，不插入空段落
        if stripped == '':
            i += 1
            continue

        # 标题（第一行非空内容，且下一行或下两行是 ====）
        if i == 0 or (i >= 1 and re.match(r'^=+$', lines[i-1].strip()) and lines[i-2].strip() == ''):
            add_title(doc, stripped)
            i += 1
            # 跳过 ==== 行
            if i < len(lines) and re.match(r'^=+$', lines[i].strip()):
                i += 1
            continue

        # 跳过分隔线
        if re.match(r'^=+$', stripped):
            i += 1
            continue

        # END 标记
        if stripped == 'END':
            add_end_marker(doc)
            i += 1
            continue

        # 章节标题 #01 #02 等
        if re.match(r'^\*\*#\d+\*\*$', stripped):
            add_section_header(doc, stripped.replace('**', ''))
            i += 1
            continue

        # 其他 bold-only 的节标题
        if re.match(r'^\*\*[^*]+话要说\*\*$', stripped):
            add_section_header(doc, stripped.replace('**', ''))
            i += 1
            continue

        # 引用块
        if stripped.startswith('>'):
            add_blockquote(doc, stripped)
            i += 1
            continue

        # 普通段落
        add_paragraph(doc, stripped)
        i += 1

    doc.save(str(docx_path))
    return True


def main():
    if len(sys.argv) < 2:
        print("用法: python convert-to-docx.py <输入.md> [输出.docx]")
        print("示例: python convert-to-docx.py 故事.md")
        print("      python convert-to-docx.py 故事.md output/故事.docx")
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"错误：文件不存在 —— {md_path}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        docx_path = Path(sys.argv[2])
    else:
        docx_path = md_path.with_suffix('.docx')

    # 确保输出目录存在
    docx_path.parent.mkdir(parents=True, exist_ok=True)

    convert_md_to_docx(md_path, docx_path)
    print(f"[OK] 已生成 Word 文档: {docx_path}")


if __name__ == '__main__':
    main()
